"""Turn an uploaded file into pages of plain text.

Every format collapses to `(page_number, text)` pairs. Only PDFs have real pages;
everything else reports a single page so citations stay meaningful.
"""

import csv
import io
import json
import logging
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

import docx
import yaml
from pypdf import PdfReader

logger = logging.getLogger(__name__)

Page = Tuple[int, str]

TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst", ".log"}
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | {
    ".pdf",
    ".docx",
    ".csv",
    ".tsv",
    ".json",
    ".jsonl",
    ".ndjson",
    ".yaml",
    ".yml",
    ".html",
    ".htm",
    ".xml",
    ".xlsx",
}

# Guard against a single enormous cell or record dominating a chunk.
MAX_CELL_CHARS = 2000
# Tabular files can hold millions of rows; cap what we ingest and say so.
MAX_TABLE_ROWS = 20000


def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_SUFFIXES


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


class _TextExtractor(HTMLParser):
    """Collect visible text, dropping script and style content."""

    SKIP = {"script", "style", "noscript", "template"}
    BREAK_AFTER = {"p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: List[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self.SKIP:
            self._skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in self.SKIP and self._skip_depth:
            self._skip_depth -= 1
        elif tag in self.BREAK_AFTER:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        stripped = data.strip()
        if stripped:
            self.parts.append(stripped + " ")

    def text(self) -> str:
        joined = "".join(self.parts)
        # Collapse the runs of blank lines the break handling leaves behind.
        lines = [line.strip() for line in joined.splitlines()]
        return "\n".join(line for line in lines if line)


def _truncate(value: str) -> str:
    value = " ".join(value.split())
    return value if len(value) <= MAX_CELL_CHARS else value[:MAX_CELL_CHARS] + "…"


def _rows_to_records(rows: Iterable[List[str]], label: Optional[str] = None) -> str:
    """Render tabular rows as `header: value` records.

    One record per line keeps a row intact through chunking, so a retrieved chunk
    carries its own column names instead of orphaned values.
    """
    iterator = iter(rows)
    try:
        header = [str(h).strip() or f"column{i + 1}" for i, h in enumerate(next(iterator))]
    except StopIteration:
        return ""

    lines: List[str] = []
    if label:
        lines.append(f"# {label}")
    lines.append("Columns: " + ", ".join(header))

    count = 0
    for row in iterator:
        if count >= MAX_TABLE_ROWS:
            lines.append(f"… truncated after {MAX_TABLE_ROWS} rows.")
            break
        cells = [
            f"{header[i] if i < len(header) else f'column{i + 1}'}: {_truncate(str(cell))}"
            for i, cell in enumerate(row)
            if str(cell).strip()
        ]
        if cells:
            lines.append(" | ".join(cells))
            count += 1
    return "\n".join(lines)


def _json_to_text(raw: str) -> str:
    """Pretty-print JSON so structure survives, falling back to the raw text."""
    try:
        parsed = json.loads(raw)
    except ValueError:
        return raw
    return json.dumps(parsed, indent=2, ensure_ascii=False, default=str)


def _jsonl_to_text(raw: str) -> str:
    records = []
    for index, line in enumerate(raw.splitlines()):
        line = line.strip()
        if not line:
            continue
        if index >= MAX_TABLE_ROWS:
            records.append(f"… truncated after {MAX_TABLE_ROWS} records.")
            break
        try:
            records.append(
                f"Record {index + 1}:\n"
                + json.dumps(json.loads(line), indent=2, ensure_ascii=False, default=str)
            )
        except ValueError:
            records.append(f"Record {index + 1}: {line}")
    return "\n\n".join(records)


def extract_pages(path: Path, original_name: str) -> List[Page]:
    suffix = Path(original_name).suffix.lower()

    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return [(i + 1, page.extract_text() or "") for i, page in enumerate(reader.pages)]

    if suffix == ".docx":
        document = docx.Document(str(path))
        blocks = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return [(1, "\n".join(blocks))]

    if suffix == ".xlsx":
        from openpyxl import load_workbook

        # read_only keeps memory flat on large workbooks.
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        sections = []
        try:
            for sheet in workbook.worksheets:
                rows = (
                    [("" if cell is None else cell) for cell in row]
                    for row in sheet.iter_rows(values_only=True)
                )
                rendered = _rows_to_records(rows, label=f"Sheet: {sheet.title}")
                if rendered:
                    sections.append(rendered)
        finally:
            workbook.close()
        return [(1, "\n\n".join(sections))]

    if suffix in {".csv", ".tsv"}:
        raw = _read_text(path)
        delimiter = "\t" if suffix == ".tsv" else None
        if delimiter is None:
            try:
                delimiter = csv.Sniffer().sniff(raw[:8192], delimiters=",;\t|").delimiter
            except csv.Error:
                delimiter = ","
        reader = csv.reader(io.StringIO(raw), delimiter=delimiter)
        return [(1, _rows_to_records(reader))]

    if suffix == ".json":
        return [(1, _json_to_text(_read_text(path)))]

    if suffix in {".jsonl", ".ndjson"}:
        return [(1, _jsonl_to_text(_read_text(path)))]

    if suffix in {".yaml", ".yml"}:
        raw = _read_text(path)
        try:
            documents = [d for d in yaml.safe_load_all(raw) if d is not None]
            return [
                (
                    1,
                    "\n\n".join(
                        yaml.safe_dump(d, sort_keys=False, allow_unicode=True)
                        for d in documents
                    ),
                )
            ]
        except yaml.YAMLError:
            return [(1, raw)]  # malformed YAML is still useful as text

    if suffix in {".html", ".htm", ".xml"}:
        parser = _TextExtractor()
        parser.feed(_read_text(path))
        return [(1, parser.text())]

    if suffix in TEXT_SUFFIXES:
        return [(1, _read_text(path))]

    raise ValueError(f"Unsupported file type: {suffix or original_name}")
