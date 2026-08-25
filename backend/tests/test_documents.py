from tests.conftest import SAMPLE_MD


def upload(client, headers, collection_id, name="notes.md", body=None, ctype="text/markdown"):
    content = SAMPLE_MD if body is None else body
    return client.post(
        f"/api/collections/{collection_id}/documents",
        headers=headers,
        files={"files": (name, content.encode("utf-8"), ctype)},
    )


def test_upload_is_ingested_and_chunked(client, auth):
    headers, collection_id, _ = auth()

    response = upload(client, headers, collection_id)
    assert response.status_code == 201, response.text

    # TestClient runs background tasks before returning, so ingest is already done.
    documents = client.get(
        f"/api/collections/{collection_id}/documents", headers=headers
    ).json()
    assert len(documents) == 1
    doc = documents[0]
    assert doc["status"] == "ready", doc.get("error")
    assert doc["chunk_count"] > 1
    assert doc["pages"] == 1
    assert doc["error"] is None

    collections = client.get("/api/collections", headers=headers).json()
    assert collections[0]["document_count"] == 1
    assert collections[0]["ready_count"] == 1


def test_multiple_file_types_supported(client, auth):
    headers, collection_id, _ = auth()
    response = client.post(
        f"/api/collections/{collection_id}/documents",
        headers=headers,
        files=[
            ("files", ("a.md", SAMPLE_MD.encode(), "text/markdown")),
            ("files", ("b.txt", SAMPLE_MD.encode(), "text/plain")),
        ],
    )
    assert response.status_code == 201, response.text
    documents = client.get(
        f"/api/collections/{collection_id}/documents", headers=headers
    ).json()
    assert {d["filename"] for d in documents} == {"a.md", "b.txt"}
    assert all(d["status"] == "ready" for d in documents)


def test_unsupported_type_rejected(client, auth):
    headers, collection_id, _ = auth()
    response = upload(client, headers, collection_id, name="virus.exe", ctype="application/octet-stream")
    assert response.status_code == 415
    assert "unsupported" in response.json()["detail"].lower()


def test_empty_file_rejected(client, auth):
    headers, collection_id, _ = auth()
    response = upload(client, headers, collection_id, name="empty.txt", body="")
    assert response.status_code == 400


def test_textless_file_marked_failed_not_crashed(client, auth):
    headers, collection_id, _ = auth()
    assert upload(client, headers, collection_id, name="blank.txt", body="   \n  \n").status_code == 201
    doc = client.get(f"/api/collections/{collection_id}/documents", headers=headers).json()[0]
    assert doc["status"] == "failed"
    assert "No extractable text" in doc["error"]


def test_delete_document_removes_its_vectors(client, auth):
    from app.services import vectorstore

    headers, collection_id, _ = auth()
    upload(client, headers, collection_id)
    doc = client.get(f"/api/collections/{collection_id}/documents", headers=headers).json()[0]
    assert vectorstore.has_index(collection_id)

    deleted = client.delete(
        f"/api/collections/{collection_id}/documents/{doc['id']}", headers=headers
    )
    assert deleted.status_code == 204
    assert client.get(
        f"/api/collections/{collection_id}/documents", headers=headers
    ).json() == []
    # Last document gone -> the whole index is dropped rather than left empty.
    assert not vectorstore.has_index(collection_id)


def test_deleting_one_document_keeps_the_others_vectors(client, auth):
    from app.services import vectorstore

    headers, collection_id, _ = auth()
    upload(client, headers, collection_id, name="keep.md")
    upload(client, headers, collection_id, name="drop.md")
    documents = client.get(
        f"/api/collections/{collection_id}/documents", headers=headers
    ).json()
    drop = next(d for d in documents if d["filename"] == "drop.md")
    keep = next(d for d in documents if d["filename"] == "keep.md")

    client.delete(f"/api/collections/{collection_id}/documents/{drop['id']}", headers=headers)

    remaining = vectorstore.retrieve(collection_id, "rehabilitation timeline")
    assert remaining, "expected the surviving document's chunks to still be searchable"
    assert {d.metadata["document_id"] for d in remaining} == {keep["id"]}


def test_documents_are_isolated_between_users(client, auth):
    headers_a, collection_a, _ = auth()
    headers_b, _, _ = auth()

    upload(client, headers_a, collection_a)

    assert client.get(f"/api/collections/{collection_a}/documents", headers=headers_b).status_code == 404
    assert client.delete(f"/api/collections/{collection_a}", headers=headers_b).status_code == 404
    assert upload(client, headers_b, collection_a).status_code == 404


def test_collection_crud(client, auth):
    headers, _, _ = auth()

    created = client.post("/api/collections", headers=headers, json={"name": "Knee papers"})
    assert created.status_code == 201
    cid = created.json()["id"]

    renamed = client.patch(f"/api/collections/{cid}", headers=headers, json={"name": "Papers"})
    assert renamed.status_code == 200 and renamed.json()["name"] == "Papers"

    assert client.delete(f"/api/collections/{cid}", headers=headers).status_code == 204
    assert cid not in [c["id"] for c in client.get("/api/collections", headers=headers).json()]


def _pdf_bytes(pages):
    """Build a real multi-page PDF so page-number metadata can be checked."""
    import io

    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    for text in pages:
        y = 750
        # Several lines per page, so the splitter has real content to chunk.
        for _ in range(12):
            pdf.drawString(72, y, text)
            y -= 18
        pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def test_pdf_pages_are_numbered_from_one(client, auth):
    from app.services import vectorstore

    headers, collection_id, _ = auth()
    body = _pdf_bytes(
        [
            "Page one covers anterior cruciate ligament anatomy in detail.",
            "Page two covers the graft options used during reconstruction.",
            "Page three covers the return to sport criteria after surgery.",
        ]
    )
    response = client.post(
        f"/api/collections/{collection_id}/documents",
        headers=headers,
        files={"files": ("acl.pdf", body, "application/pdf")},
    )
    assert response.status_code == 201, response.text

    doc = client.get(f"/api/collections/{collection_id}/documents", headers=headers).json()[0]
    assert doc["status"] == "ready", doc.get("error")
    assert doc["pages"] == 3
    assert doc["chunk_count"] >= 3

    hits = vectorstore.retrieve(collection_id, "graft options during reconstruction")
    assert hits
    pages = {h.metadata["page"] for h in hits}
    assert pages <= {1, 2, 3} and min(pages) >= 1, f"1-indexed pages expected, got {pages}"
    assert all(h.metadata["total_pages"] == 3 for h in hits)
    assert all(h.metadata["filename"] == "acl.pdf" for h in hits)


def test_docx_is_extracted(client, auth):
    import io

    import docx

    headers, collection_id, _ = auth()
    document = docx.Document()
    for _ in range(8):
        document.add_paragraph(
            "Rehabilitation after ACL reconstruction progresses through distinct phases."
        )
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Phase"
    table.rows[0].cells[1].text = "Timeline"
    buffer = io.BytesIO()
    document.save(buffer)

    response = client.post(
        f"/api/collections/{collection_id}/documents",
        headers=headers,
        files={
            "files": (
                "plan.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201, response.text
    doc = client.get(f"/api/collections/{collection_id}/documents", headers=headers).json()[0]
    assert doc["status"] == "ready", doc.get("error")
    assert doc["chunk_count"] >= 1
